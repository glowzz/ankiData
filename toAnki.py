# -*- coding: utf-8 -*-
"""
Created on Sun May 12 11:40:55 2019

@author: Administrator
"""
#import numpy as np
import pandas as pd
import os
import shutil
import datetime
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def getExcData(Exc_path,sheetNum):
    ExcData=[]
    df1 = pd.read_excel(Exc_path,sheet_name=sheetNum)
    df2 = df1.dropna(thresh=2)
    column_name=df2.columns[1]
    for index in df2.index:
        yearErr=df2.loc[index].dropna().values
        yearE=yearErr[0]
        yearQnum=yearErr[1:]
    #print(yearE)
        for Enum in yearQnum:
            QpicName="%02d.png" % Enum
            ApicName="%02dAN.png" % Enum
            ExcData.append([yearE,QpicName,ApicName])
    return ExcData,column_name
    
    
def formAnkiData(fromData):
    ankiData=[]
    for Enum in fromData:
        yearE= Enum[0]
        Qu_picName= Enum[1]
        An_picName= Enum[2]
        outName=yearE+Qu_picName
       
        str1="\"<div>%s</div><img src=\"\"%s\"\" />\"" % (outName[:-4],outName)
        str2="\"<img src=\"\"%s\"\" />\"" % (yearE+An_picName)
        ankiData.append([str1,yearE,str2])
    return ankiData
    



def copyfile(from_path,to_path,filedata,index=1):
    if index==1:
        from_path=os.path.join(from_path,'QU')
    else:
        from_path=os.path.join(from_path,'AN')
        
    if not os.path.exists(to_path):
        os.mkdir(to_path)
#print(os.path.join(from_path,filename1))
        
    for eachfile in filedata:
        file_path=os.path.join(from_path,eachfile[0],eachfile[1])
        if os.path.exists(file_path):
            shutil.copy(file_path,os.path.join(to_path,eachfile[0]+eachfile[index]))
        else:
            print(file_path+" no found")            
    #shutil.copy(os.path.join(from_path,fileName),os.path.join(to_path,asName))
    
def writeAnki_txt(to_filepath,data2write):
    with open(to_filepath, 'w') as f:
        for a in data2write:
            f.write(a[0]+"\t"+a[1]+"\t"+a[2]+"\t\n")#for improve


def writeDoc(Save_filepath,from_path,filedata):
    document = Document()                #以默认模板建立文档对象

    document.add_paragraph()
    p_detail = document.add_paragraph()
    r_detail = p_detail.add_run(u'真题错题集')
    r_detail.font.bold = True
    p_detail.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    from_path1=os.path.join(from_path,'QU')
    for eachfile in filedata:
        document.add_paragraph(eachfile[0]+eachfile[1].split('.')[0])
        file_path=os.path.join(from_path1,eachfile[0],eachfile[1])
        if os.path.exists(file_path):
            document.add_picture(file_path,width=Inches(6))
        else:
            print(file_path+" no found")
        document.add_paragraph()
        document.add_paragraph()
        document.add_paragraph()
        

    p_detail = document.add_paragraph()
    r_detail = p_detail.add_run(u'答案')
    r_detail.font.bold = True
    p_detail.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    from_path2=os.path.join(from_path,'AN')
    
    for eachfile in filedata:
        document.add_paragraph(eachfile[0]+eachfile[2].split('.')[0])
        
        file_path=os.path.join(from_path2,eachfile[0],eachfile[1])
        if os.path.exists(file_path):
            document.add_picture(file_path,width=Inches(4))
        else:
            print(file_path+" no found")
            
    document.save(Save_filepath)  

def copy_all_file(from_path,to_path,filedata):
    #lists = os.listdir(from_path)
    for eachfile in filedata:
        #print(i)
        file_path=os.path.join(from_path,eachfile[0]+eachfile[1])
        if os.path.exists(file_path):
            shutil.copy(file_path,to_path)
        else:
            print(file_path+" no found") 
            
        file_path=os.path.join(from_path,eachfile[0]+eachfile[2])
        if os.path.exists(file_path):
            shutil.copy(file_path,to_path)
        else:
            print(file_path+" no found") 
             
            
            
'''           
        shutil.copy(os.path.join(from_path,eachfile[0]+eachfile[1]),to_path)
        
        shutil.copy(os.path.join(from_path,eachfile[0]+eachfile[2]),to_path)  
'''        
    
def main():
    
    now_time = datetime.datetime.now()
    detester =now_time.strftime('%Y-%m-%d')
    from_path="E:\\Users\\Administrator\\Documents\\Snagit"
    to_path=r"C:\Users\Admin\AppData\Roaming\Anki2\000\collection.media"
    Exc_path="G:\\新建文件夹\\错题to_anki"
    Exc_filename="错题all_other1.xlsx"
    for i in range(4):
        
        ankiData=[]
        ExcData=[]
        ExcData,student_name=getExcData(os.path.join(Exc_path,Exc_filename),i)
        
        ankiData=formAnkiData(ExcData)  
    #    print(ExcData)
    
        print(student_name)
        writeDoc(os.path.join(Exc_path,student_name+"错题"+detester+".docx"),from_path,ExcData)
       
        writeAnki_txt(os.path.join(Exc_path,student_name+detester+".txt"),ankiData)
        copyfile(from_path,os.path.join(Exc_path,"media"),ExcData,1)
        copyfile(from_path,os.path.join(Exc_path,"media"),ExcData,2)
        copy_all_file(os.path.join(Exc_path,"media"),to_path,ExcData)
       

     
            
    
    #print(ankiData)
 #   frame = pd.DataFrame(ankiData)
 #   frame.to_csv(os.path.join(to_path,"out.txt"))
    
main()
        
        
